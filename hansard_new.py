# import os
# import tempfile
# import streamlit as st
# import streamlit_authenticator as stauth
# from langchain_community.document_loaders import PyPDFLoader
# from langchain_openai import OpenAIEmbeddings, ChatOpenAI
# from langchain_core.output_parsers import StrOutputParser
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_core.runnables import RunnableParallel, RunnablePassthrough
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import FAISS
# from docx import Document as DocxDocument
# from io import BytesIO
# import yaml
# from yaml.loader import SafeLoader

# # Page config
# st.set_page_config(page_title="Hansard Analyzer", layout="wide")

# # Initialize session state
# if 'authentication_status' not in st.session_state:
#     st.session_state['authentication_status'] = None
# if 'name' not in st.session_state:
#     st.session_state['name'] = None
# if 'username' not in st.session_state:
#     st.session_state['username'] = None

# def load_config():
#     """Load configuration from secrets"""
#     try:
#         return {
#             'credentials': yaml.safe_load(st.secrets["general"]["credentials"]),
#             'cookie_name': st.secrets["general"]["cookie_name"],
#             'cookie_key': st.secrets["general"]["cookie_key"],
#             'cookie_expiry_days': st.secrets["general"]["cookie_expiry_days"],
#             'openai_api_key': st.secrets["general"]["OPENAI_API_KEY"]
#         }
#     except Exception as e:
#         st.error(f"Error loading configuration: {str(e)}")
#         st.stop()

# @st.cache_data(show_spinner=False)
# def process_documents(openai_api_key, model_name, uploaded_files, query, prompt_type):
#     """Process documents and generate analysis"""
#     try:
#         # Initialize progress indicators
#         progress_text = st.empty()
#         progress_text.text("Loading documents...")
#         loading_progress = st.progress(0)
#         processing_progress = st.progress(0)

#         # Initialize language models
#         embeddings = OpenAIEmbeddings(
#             model='text-embedding-3-small',
#             openai_api_key=openai_api_key
#         )
#         llm = ChatOpenAI(
#             temperature=0,
#             model_name=model_name,
#             max_tokens=4000,
#             openai_api_key=openai_api_key
#         )

#         # Process documents
#         docs = []
#         total_files = len(uploaded_files)
        
#         for i, uploaded_file in enumerate(uploaded_files):
#             progress_text.text(f"Processing file {i+1} of {total_files}...")
#             with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#                 tmp_file.write(uploaded_file.read())
#                 tmp_file_path = tmp_file.name

#             loader = PyPDFLoader(file_path=tmp_file_path)
#             docs.extend(loader.load())
            
#             os.remove(tmp_file_path)
#             loading_progress.progress((i + 1) / total_files)

#         progress_text.text("Splitting documents...")
#         text_splitter = RecursiveCharacterTextSplitter(
#             chunk_size=1500,
#             chunk_overlap=300
#         )
#         splits = text_splitter.split_documents(docs)
#         processing_progress.progress(0.4)

#         progress_text.text("Creating vector store...")
#         vectorstore = FAISS.from_documents(splits, embeddings)
#         retriever = vectorstore.as_retriever(
#             search_kwargs={"k": 5}
#         )
#         processing_progress.progress(0.6)

#         progress_text.text("Analyzing content...")
        
#         # Use the prompts from session state
#         if prompt_type == "Default":
#             prompt = ChatPromptTemplate.from_template(st.session_state.default_prompt)
#         else:
#             prompt = ChatPromptTemplate.from_template(st.session_state.detailed_prompt)

#         chain = (
#             RunnableParallel(
#                 {"context": retriever, "input": RunnablePassthrough()}
#             )
#             | prompt
#             | llm
#             | StrOutputParser()
#         )
        
#         processing_progress.progress(0.8)
#         response = chain.invoke(query)
#         processing_progress.progress(1.0)

#         # Create document
#         buffer = BytesIO()
#         doc = DocxDocument()
#         doc.add_paragraph(f"Question: {query}\n\nAnswer:\n")
#         doc.add_paragraph(response)
#         doc.save(buffer)
#         buffer.seek(0)

#         # Clear progress indicators
#         progress_text.empty()
#         loading_progress.empty()
#         processing_progress.empty()

#         return response, buffer

#     except Exception as e:
#         st.error(f"Error in document processing: {str(e)}")
#         return None, None

# def main():
#     config = load_config()
    
#     # Create authenticator
#     authenticator = stauth.Authenticate(
#         config['credentials'],
#         config['cookie_name'],
#         config['cookie_key'],
#         config['cookie_expiry_days']
#     )

#     st.title("Hansard Analyzer")

#     # Handle Authentication
#     if not st.session_state['authentication_status']:
#         authentication_status, name, username = authenticator.login()
        
#         if authentication_status:
#             st.session_state['authentication_status'] = authentication_status
#             st.session_state['name'] = name
#             st.session_state['username'] = username
#             st.rerun()
#         elif authentication_status is False:
#             st.error("Username/password is incorrect")
#         elif authentication_status is None:
#             st.warning("Please enter your username and password")
#         return


#     # Main application
#     if st.session_state['authentication_status']:
#         # Initialize session state for prompts if not exists
#         if 'default_prompt' not in st.session_state:
#             st.session_state.default_prompt = """You are provided with a context extracted from Canadian parliamentary debates (Hansard) concerning various political issues.
# Answer the question by focusing on the relevant party based on the question. Provide the five to six main points and conclusion.
# {context}
# Question: {input}"""

#         if 'custom_prompt' not in st.session_state:
#             st.session_state.custom_prompt = """You are provided with a context extracted from Canadian parliamentary debates (Hansard) concerning various political issues. 
# Answer the question by focusing on the relevant party based on the question. Provide the five to six main points and conclusion. 

# If the question asks about the Liberal Party, focus on the Liberal Party's viewpoint. If the question asks about the Conservative Party, focus on the Conservative Party's viewpoint. 

# Provide detailed information including their proposals, policy stance, and any arguments made during the debate.

# <context>
# {context}
# </context>

# Question: {input}

# **Main Points:**
# 1- Six main points summarizing the party's stance 

# **Supporting Quotes:**
# 2- List specific quotes that support the analysis, including the names of the individuals who made them or references from the debates

# **Potential Implications of Each Party's Stance:**
# 3 - Any significant points raised during the debate, including potential implications of each party's stance for each Question

# **Conclusion:**
# 4- Summarize the party's stance and its implications"""

#         # Sidebar
#         with st.sidebar:
#             st.write(f"Welcome *{st.session_state['name']}*")
            
#             # Logout button
#             if st.button("Logout", key='logout_button'):
#                 # Clear all session states
#                 for key in list(st.session_state.keys()):
#                     del st.session_state[key]
#                 st.rerun()
            
#             st.markdown("### Settings")
            
#             model_name = st.selectbox(
#                 "Select Model",
#                 ["gpt-4o-mini", "gpt-4o"]
#             )
            
#             # Add prompt selection
#             prompt_type = st.radio(
#                 "Select Prompt Type",
#                 ["Default", "Custom"],
#                 help="Choose between default or custom analysis prompts"
#             )
            
#             # Show current prompt based on selection
#             with st.expander("View Current Prompt", expanded=False):
#                 if prompt_type == "Default":
#                     st.code(st.session_state.default_prompt)
#                 else:  # Custom
#                     st.code(st.session_state.custom_prompt)
            
#             # Expandable and editable custom prompt section
#             if prompt_type == "Custom":
#                 with st.expander("📝 Edit Custom Prompt", expanded=True):
#                     # CSS to make the text area stretchable
#                     st.markdown("""
#                         <style>
#                             .stTextArea textarea {
#                                 height: 400px !important;
#                                 min-width: 100% !important;
#                                 resize: both !important;
#                                 overflow: auto !important;
#                             }
#                         </style>
#                         """, unsafe_allow_html=True)
                    
#                     custom_prompt = st.text_area(
#                         "Edit Custom Prompt",
#                         value=st.session_state.custom_prompt,
#                         height=400,
#                         key="custom_prompt_editor",
#                         label_visibility="collapsed"  # This hides the label for cleaner look
#                     )
                    
#                     col1, col2 = st.columns(2)
#                     with col1:
#                         if st.button("Save Custom Prompt"):
#                             st.session_state.custom_prompt = custom_prompt
#                             st.success("Custom prompt updated!")
                    
#                     with col2:
#                         if st.button("Reset Custom Prompt"):
#                             st.session_state.custom_prompt = st.session_state.custom_prompt
#                             st.success("Custom prompt reset!")
            
#             uploaded_files = st.file_uploader(
#                 "Upload PDF files",
#                 type="pdf",
#                 accept_multiple_files=True
#             )

#         # Main content
#         query = st.text_input(
#             "Enter your query",
#             value="What is the position of the Liberal Party on Carbon Pricing?"
#         )

#         col1, col2 = st.columns([3, 1])
#         with col1:
#             analyze_button = st.button("Analyze", type="primary")
#         with col2:
#             clear_button = st.button("Clear Results")

#         if clear_button:
#             st.session_state.analysis_result = None
#             st.rerun()

#         if analyze_button:
#             if uploaded_files and query:
#                 with st.spinner("Analyzing documents..."):
#                     answer, buffer = process_documents(
#                         config['openai_api_key'],
#                         model_name,
#                         uploaded_files,
#                         query,
#                         prompt_type
#                     )
#                     if answer and buffer:
#                         # Store results in session state
#                         st.session_state.analysis_result = {
#                             'answer': answer,
#                             'buffer': buffer,
#                             'query': query
#                         }
#                         st.rerun()
#             else:
#                 st.warning("Please upload PDF files and enter a query.")

#         # Display results if available
#         if 'analysis_result' in st.session_state and st.session_state.analysis_result:
#             st.markdown("### Analysis Results")
#             with st.container():
#                 st.markdown(f"**Question:** {st.session_state.analysis_result['query']}")
#                 st.markdown("**Answer:**")
#                 st.markdown(st.session_state.analysis_result['answer'])
                
#                 st.download_button(
#                     label="📥 Download Analysis as Word Document",
#                     data=st.session_state.analysis_result['buffer'],
#                     file_name="hansard_analysis.docx",
#                     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                 )

        
# if __name__ == "__main__":
#     main()


import os
import tempfile
import streamlit as st
import streamlit_authenticator as stauth
from langchain_community.document_loaders import PyPDFLoader
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableParallel, RunnablePassthrough
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from docx import Document as DocxDocument
from io import BytesIO
import yaml
from yaml.loader import SafeLoader

# Page config
st.set_page_config(page_title="Hansard Analyzer", layout="wide")

# Initialize session state
if 'authentication_status' not in st.session_state:
    st.session_state['authentication_status'] = None
if 'name' not in st.session_state:
    st.session_state['name'] = None
if 'username' not in st.session_state:
    st.session_state['username'] = None


def load_config():
    """Load configuration from secrets"""
    try:
        return {
            'credentials': yaml.safe_load(st.secrets["general"]["credentials"]),
            'cookie_name': st.secrets["general"]["cookie_name"],
            'cookie_key': st.secrets["general"]["cookie_key"],
            'cookie_expiry_days': st.secrets["general"]["cookie_expiry_days"],
            'openai_api_key': st.secrets["general"]["OPENAI_API_KEY"]
        }
    except Exception as e:
        st.error(f"Error loading configuration: {str(e)}")
        st.stop()

@st.cache_data(show_spinner=False)
def process_documents(openai_api_key, model_name, uploaded_files, query, prompt_type):
    """Process documents and generate analysis"""
    try:
        # Initialize progress indicators
        progress_text = st.empty()
        progress_text.text("Loading documents...")
        loading_progress = st.progress(0)
        processing_progress = st.progress(0)

        # Initialize language models
        embeddings = OpenAIEmbeddings(
            model='text-embedding-3-small',
            openai_api_key=openai_api_key
        )
        llm = ChatOpenAI(
            temperature=0,
            model_name=model_name,
            max_tokens=4000,
            openai_api_key=openai_api_key
        )

        # Process documents
        docs = []
        total_files = len(uploaded_files)
        
        for i, uploaded_file in enumerate(uploaded_files):
            progress_text.text(f"Processing file {i+1} of {total_files}...")
            with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_file_path = tmp_file.name

            loader = PyPDFLoader(file_path=tmp_file_path)
            docs.extend(loader.load())
            
            os.remove(tmp_file_path)
            loading_progress.progress((i + 1) / total_files)

        progress_text.text("Splitting documents...")
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=300
        )
        splits = text_splitter.split_documents(docs)
        processing_progress.progress(0.4)

        progress_text.text("Creating vector store...")
        vectorstore = FAISS.from_documents(splits, embeddings)
        retriever = vectorstore.as_retriever(
            search_kwargs={"k": 5}
        )
        processing_progress.progress(0.6)

        progress_text.text("Analyzing content...")
        
        # Use the prompts from session state
        if prompt_type == "Default":
            prompt = ChatPromptTemplate.from_template(st.session_state.default_prompt)
        else:
            
            prompt = ChatPromptTemplate.from_template(st.session_state.custom_prompt)

        chain = (
            RunnableParallel(
                {"context": retriever, "input": RunnablePassthrough()}
            )
            | prompt
            | llm
            | StrOutputParser()
        )
        
        processing_progress.progress(0.8)
        response = chain.invoke(query)
        processing_progress.progress(1.0)

        # Create document
        buffer = BytesIO()
        doc = DocxDocument()
        doc.add_paragraph(f"Question: {query}\n\nAnswer:\n")
        doc.add_paragraph(response)
        doc.save(buffer)
        buffer.seek(0)

        # Clear progress indicators
        progress_text.empty()
        loading_progress.empty()
        processing_progress.empty()

        return response, buffer

    except Exception as e:
        st.error(f"Error in document processing: {str(e)}")
        return None, None

def main():
    config = load_config()
    
    # Create authenticator
    authenticator = stauth.Authenticate(
        config['credentials'],
        config['cookie_name'],
        config['cookie_key'],
        config['cookie_expiry_days']
    )

    st.title("Hansard Analyzer")

    # Handle Authentication
    if st.session_state['authentication_status'] is None:
        # Center the login form using columns
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            login_result = authenticator.login()
            
            if login_result is not None:
                authentication_status, name, username = login_result
                
                if authentication_status:
                    # Successful login
                    st.session_state['authentication_status'] = True
                    st.session_state['name'] = name
                    st.session_state['username'] = username
                    st.success(f"Welcome, {name}!")
                    st.experimental_rerun()
                elif authentication_status is False:
                    # Incorrect credentials
                    st.error("Username/password is incorrect")
                elif authentication_status is None:
                    # No credentials provided
                    st.warning("Please enter your username and password")
        return
    # Main application logic after successful login
    if st.session_state['authentication_status']:
        # Initialize session state for prompts if not exists
        if 'default_prompt' not in st.session_state:
            st.session_state.default_prompt = """You are provided with a context extracted from Canadian parliamentary debates (Hansard) concerning various political issues.
                Answer the question by focusing on the relevant party based on the question. Provide the five to six main points and conclusion.
                {context}
                Question: {input}"""

        if 'custom_prompt' not in st.session_state:
            st.session_state.custom_prompt = """You are provided with a context extracted from Canadian parliamentary debates (Hansard) concerning various political issues. 
                Answer the question by focusing on the relevant party based on the question. Provide the five to six main points and conclusion. 

                If the question asks about the Liberal Party, focus on the Liberal Party's viewpoint. If the question asks about the Conservative Party, focus on the Conservative Party's viewpoint. 

                Provide detailed information including their proposals, policy stance, and any arguments made during the debate.

                <context>
                {context}
                </context>

                Question: {input}

                **Main Points:**
                1- Six main points summarizing the party's stance 

                **Supporting Quotes:**
                2- List specific quotes that support the analysis, including the names of the individuals who made them or references from the debates

                **Potential Implications of Each Party's Stance:**
                3 - Any significant points raised during the debate, including potential implications of each party's stance for each Question

                **Conclusion:**
                4- Summarize the party's stance and its implications"""

        # Sidebar
        with st.sidebar:
            st.write(f"Welcome *{st.session_state['name']}*")
            
            # Logout button
            authenticator.logout('Logout', 'sidebar')
            
            st.markdown("### Settings")
            
            model_name = st.selectbox(
                "Select Model",
                ["gpt-4o-mini", "gpt-4o"]
            )
            
            # Add prompt selection
            prompt_type = st.radio(
                "Select Prompt Type",
                ["Default", "Custom"],
                help="Choose between default or custom analysis prompts"
            )
            
            # Show current prompt based on selection
            with st.expander("View Current Prompt", expanded=False):
                if prompt_type == "Default":
                    st.code(st.session_state.default_prompt)
                else:  # Custom
                    st.code(st.session_state.custom_prompt)
            
            # Expandable and editable custom prompt section
            if prompt_type == "Custom":
                with st.expander("📝 Edit Custom Prompt", expanded=True):
                    # CSS to make the text area stretchable
                    st.markdown("""
                        <style>
                            .stTextArea textarea {
                                height: 400px !important;
                                min-width: 100% !important;
                                resize: both !important;
                                overflow: auto !important;
                            }
                        </style>
                        """, unsafe_allow_html=True)
                    
                    custom_prompt = st.text_area(
                        "Edit Custom Prompt",
                        value=st.session_state.custom_prompt,
                        height=400,
                        key="custom_prompt_editor",
                        label_visibility="collapsed"  # This hides the label for cleaner look
                    )
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("Save Custom Prompt"):
                            st.session_state.custom_prompt = custom_prompt
                            st.success("Custom prompt updated!")
                    
                    with col2:
                        if st.button("Reset Custom Prompt"):
                            st.session_state.custom_prompt = st.session_state.custom_prompt
                            st.success("Custom prompt reset!")
            
            uploaded_files = st.file_uploader(
                "Upload PDF files",
                type="pdf",
                accept_multiple_files=True
            )

        # Main content
        query = st.text_input(
            "Enter your query",
            value="What is the position of the Liberal Party on Carbon Pricing?"
        )

        col1, col2 = st.columns([3, 1])
        with col1:
            analyze_button = st.button("Analyze", type="primary")
        with col2:
            clear_button = st.button("Clear Results")

        if clear_button:
            st.session_state.analysis_result = None
            st.rerun()

        if analyze_button:
            if uploaded_files and query:
                with st.spinner("Analyzing documents..."):
                    answer, buffer = process_documents(
                        config['openai_api_key'],
                        model_name,
                        uploaded_files,
                        query,
                        prompt_type
                    )
                    if answer and buffer:
                        # Store results in session state
                        st.session_state.analysis_result = {
                            'answer': answer,
                            'buffer': buffer,
                            'query': query
                        }
                        st.rerun()
            else:
                st.warning("Please upload PDF files and enter a query.")

        # Display results if available
        if 'analysis_result' in st.session_state and st.session_state.analysis_result:
            st.markdown("### Analysis Results")
            with st.container():
                st.markdown(f"**Question:** {st.session_state.analysis_result['query']}")
                st.markdown("**Answer:**")
                st.markdown(st.session_state.analysis_result['answer'])
                
                st.download_button(
                    label="📥 Download Analysis as Word Document",
                    data=st.session_state.analysis_result['buffer'],
                    file_name="hansard_analysis.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
